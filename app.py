from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import requests
import json
import os
import time
from dotenv import load_dotenv
from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uuid

# 加载环境变量
load_dotenv()

app = FastAPI(title="智能文档生成器")

# 配置静态文件和模板
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

# 创建必要的目录
Path("templates").mkdir(exist_ok=True)
Path("static").mkdir(exist_ok=True)
Path("downloads").mkdir(exist_ok=True)

# 大模型API配置
MODEL_API_URL = os.getenv("MODEL_API_URL", "")
MODEL_API_KEY = os.getenv("MODEL_API_KEY", "")  # 默认使用演示密钥，建议在生产环境中使用环境变量
MODEL_NAME = os.getenv("MODEL_NAME", "")
# 存储进行中的文档生成任务
active_generations = {}

# 在文件开头定义一个清理文本的函数，移除Markdown特殊符号
def clean_markdown_text(text):
    """移除Markdown格式中的特殊符号，如*、`等"""
    # 移除*符号（同时保留文本）
    text = text.replace('*', '')
    # 移除其他可能的Markdown格式符号
    text = text.replace('`', '')
    text = text.replace('>', '')
    # 处理其他可能的格式符号
    text = text.replace('<br>', '')
    text = text.replace('<br/>', '')
    text = text.replace('&nbsp;', ' ')
    # 移除多余空格
    while '  ' in text:
        text = text.replace('  ', ' ')
    return text.strip()

# 添加标题处理函数，确保标题格式一致且不重复
def process_document_headings(content):
    """处理文档中的标题，确保格式一致、层级合理且不重复"""
    lines = content.split('\n')
    processed_lines = []
    seen_headings = {}  # 用于记录已经出现过的标题
    current_level = 0
    last_heading_level = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            processed_lines.append('')
            continue
        
        # 检测是否是标题行
        heading_level = 0
        for i in range(4, 0, -1):  # 从#### 到 # 依次检查
            prefix = '#' * i + ' '
            if line.startswith(prefix):
                heading_level = i
                title_text = line[len(prefix):].strip()
                title_text = clean_markdown_text(title_text)
                
                # 标题内容清理和规范化
                title_text = title_text.strip('.,;:')  # 移除可能的结尾标点
                
                # 检查是否是重复或极度相似的标题
                lower_title = title_text.lower()
                if lower_title in seen_headings:
                    # 如果完全相同的标题已经存在，跳过这个标题
                    if seen_headings[lower_title] == heading_level:
                        continue
                
                # 确保标题层级合理（不能从一级跳到三级）
                if heading_level > last_heading_level + 1 and last_heading_level > 0:
                    heading_level = last_heading_level + 1
                
                # 记录该标题
                seen_headings[lower_title] = heading_level
                last_heading_level = heading_level
                
                # 更新为正确的标记和格式
                line = '#' * heading_level + ' ' + title_text
                break
        
        processed_lines.append(line)
    
    return '\n'.join(processed_lines)

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/generate-outline")
async def generate_outline(topic: str = Form(), key_points: str = Form(None)):
    """根据主题生成大纲"""
    if not topic:
        raise HTTPException(status_code=400, detail="主题不能为空")
    
    # 构建发送到大模型API的请求
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MODEL_API_KEY}"
    }
    
    # 构建提示，明确要求大模型生成大纲
    prompt = f"请为主题'{topic}'生成一个详细的、结构化的大纲。大纲应包含章节标题和子章节，以及每个部分的简要说明。请确保大纲逻辑清晰，结构合理。格式要求：1.使用Markdown格式编写，2.使用#、##、###表示不同层级的标题，3.每个子章节可以用1-2句话描述其内容要点。"
    
    # 如果有重点提要，加入提示
    if key_points and key_points.strip():
        prompt += f"\n\n请务必确保大纲围绕以下关键点展开：\n{key_points}\n\n这些要点必须体现在大纲的相应章节中。"
    
    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是一个专业的大纲生成助手，擅长创建结构化、有逻辑性的大纲。"},
            {"role": "user", "content": prompt}
        ]
    }
    
    try:
        # 发送请求到大模型API
        response = requests.post(MODEL_API_URL, headers=headers, json=data)
        response.raise_for_status()  # 检查是否请求成功
        
        # 解析响应
        result = response.json()
        outline_text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not outline_text:
            raise HTTPException(status_code=500, detail="大模型返回内容为空")
        
        return {"outline": outline_text}
    
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"API请求失败: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理大纲生成时出错: {str(e)}")


@app.post("/generate-fulltext-section")
async def generate_fulltext_section(outline_section: str = Form(), context: str = Form("")):
    """根据大纲部分生成对应的全文内容"""
    if not outline_section:
        raise HTTPException(status_code=400, detail="大纲内容不能为空")
    
    # 构建发送到大模型API的请求
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MODEL_API_KEY}"
    }
    
    # 构建提示，要求大模型根据大纲部分生成详细内容
    prompt = f"我需要你根据以下大纲部分生成详细的文章内容。请写出完整、连贯、内容丰富的段落，要有论述和解释，不要只是简单地重复大纲内容。\n\n大纲部分:\n{outline_section}\n\n"
    
    # 如果有上下文，则加入提示
    if context:
        prompt += f"\n已有内容作为参考上下文:\n{context}\n\n请确保生成的内容与上下文保持一致性和连贯性。"
    
    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是一位专业的内容撰写专家，擅长将大纲扩展为详细的文章内容。生成的内容应该结构清晰、论述充分、内容翔实。"},
            {"role": "user", "content": prompt}
        ]
    }
    
    try:
        # 发送请求到大模型API
        response = requests.post(MODEL_API_URL, headers=headers, json=data)
        response.raise_for_status()  # 检查是否请求成功
        
        # 解析响应
        result = response.json()
        section_text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not section_text:
            raise HTTPException(status_code=500, detail="大模型返回内容为空")
        
        return {"section_text": section_text}
    
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"API请求失败: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理全文生成时出错: {str(e)}")


@app.post("/start-fulltext-generation")
async def start_fulltext_generation(outline: str = Form(), key_points: str = Form(None)):
    """开始全文生成任务"""
    if not outline:
        raise HTTPException(status_code=400, detail="大纲内容不能为空")
    
    # 生成唯一任务ID
    task_id = str(uuid.uuid4())
    
    # 初始化任务
    active_generations[task_id] = {
        "outline": outline,
        "key_points": key_points or "",  # 存储重点提要
        "status": "starting",
        "progress": 0,
        "full_text": "",
        "start_time": time.time(),
        "sections_total": 0,
        "sections_completed": 0
    }
    
    # 解析大纲，计算总部分数
    sections = []
    lines = outline.split("\n")
    
    current_section = ""
    for line in lines:
        if line.strip() and (line.startswith("#") or current_section):
            if line.startswith("#"):
                if current_section:
                    sections.append(current_section)
                current_section = line
            else:
                current_section += "\n" + line
    
    # 添加最后一个部分
    if current_section:
        sections.append(current_section)
    
    active_generations[task_id]["sections"] = sections
    active_generations[task_id]["sections_total"] = len(sections)
    active_generations[task_id]["status"] = "in_progress"
    
    return {"task_id": task_id, "total_sections": len(sections)}


@app.get("/generation-status/{task_id}")
async def get_generation_status(task_id: str):
    """获取文档生成进度"""
    if task_id not in active_generations:
        raise HTTPException(status_code=404, detail="生成任务不存在")
    
    task = active_generations[task_id]
    return {
        "status": task["status"],
        "progress": task["progress"],
        "sections_completed": task["sections_completed"],
        "sections_total": task["sections_total"]
    }


@app.post("/process-next-section")
async def process_next_section(task_id: str = Form()):
    """处理下一个大纲部分"""
    if task_id not in active_generations:
        raise HTTPException(status_code=404, detail="生成任务不存在")
    
    task = active_generations[task_id]
    
    # 检查是否已完成所有部分
    if task["sections_completed"] >= task["sections_total"]:
        return {
            "completed": True, 
            "progress": 100,
            "status": "completed",
            "full_text": task["full_text"]
        }
    
    # 获取下一个待处理部分
    section_index = task["sections_completed"]
    current_section = task["sections"][section_index]
    
    try:
        # 调用API生成该部分内容
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {MODEL_API_KEY}"
        }
        
        # 提取部分标题作为提示
        section_title = current_section.split("\n")[0] if "\n" in current_section else current_section
        
        # 构建提示
        prompt = f"请根据以下大纲部分生成详细的文章内容。生成的内容应包含完整的段落，有充分的论述、解释和例子，而不仅仅是重复大纲内容。\n\n大纲部分:\n{current_section}\n\n"
        
        # 添加重点提要以确保内容不偏题
        if task.get("key_points") and task["key_points"].strip():
            prompt += f"\n重要：请确保生成的内容与以下关键点/核心要求相符合：\n{task['key_points']}\n\n请将这些要点自然地融入到内容中，确保内容紧扣主题而不偏离。\n\n"
        
        # 添加上下文以保持连贯性
        if task["full_text"]:
            # 提取最后一部分作为上下文
            last_part = task["full_text"].split("\n\n")[-3:] if "\n\n" in task["full_text"] else task["full_text"]
            last_part = "\n\n".join(last_part) if isinstance(last_part, list) else last_part
            prompt += f"已有内容作为上下文参考:\n{last_part}\n\n请确保新内容与已有内容保持连贯性。"
        
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": "你是一位专业的内容撰写专家，擅长将大纲扩展为详细的文章内容。生成的内容应该结构清晰、论述充分、内容翔实。你会确保内容紧扣主题要点，不偏离核心内容。"},
                {"role": "user", "content": prompt}
            ]
        }
        
        # 发送请求
        response = requests.post(MODEL_API_URL, headers=headers, json=data)
        response.raise_for_status()
        
        # 解析响应
        result = response.json()
        section_text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not section_text:
            section_text = "系统无法生成此部分内容，请稍后重试。"
        
        # 更新全文
        heading_level = section_title.count("#")
        formatted_title = section_title.replace("#", "").strip()
        
        # 格式化标题
        if heading_level == 1:
            formatted_section = f"\n\n## {formatted_title}\n\n{section_text}"
        elif heading_level == 2:
            formatted_section = f"\n\n### {formatted_title}\n\n{section_text}"
        else:
            formatted_section = f"\n\n#### {formatted_title}\n\n{section_text}"
        
        # 添加到全文
        task["full_text"] += formatted_section
        
        # 更新进度
        task["sections_completed"] += 1
        task["progress"] = int((task["sections_completed"] / task["sections_total"]) * 100)
        
        return {
            "completed": task["sections_completed"] >= task["sections_total"],
            "progress": task["progress"],
            "status": "completed" if task["sections_completed"] >= task["sections_total"] else "in_progress",
            "section_text": formatted_section,
            "current_section": current_section
        }
        
    except Exception as e:
        # 出错时，将状态设为错误
        task["status"] = "error"
        raise HTTPException(status_code=500, detail=f"处理部分生成时出错: {str(e)}")


@app.post("/export-to-doc")
async def export_to_doc(content: str = Form(), content_type: str = Form("outline")):
    """将内容导出为Word文档"""
    if not content:
        raise HTTPException(status_code=400, detail="内容不能为空")
    
    try:
        # 预处理内容，确保标题格式一致且不重复
        content = process_document_headings(content)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档基本属性
        core_properties = doc.core_properties
        core_properties.title = "智能文档生成器" + ("大纲" if content_type == "outline" else "完整文档")
        core_properties.language = "zh-CN"
        
        # 设置页面边距和纸张大小
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
            section.page_height = Inches(11.69)  # A4纸高度
            section.page_width = Inches(8.27)    # A4纸宽度
        
        # 获取主标题（从内容的第一行提取，如果有#号）
        main_title = "生成的文档"
        if content and "\n" in content:
            first_line = content.split('\n')[0].strip()
            if first_line.startswith("# "):
                main_title = clean_markdown_text(first_line[2:].strip())
        
        # 添加封面
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  # 添加封面分页符
        p = doc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n\n\n")  # 添加空行
        
        # 封面使用从内容提取的标题，而不是固定文本
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(main_title)
        run.font.size = Pt(26)
        run.font.name = '黑体'
        run.font.bold = True
        
        # 日期
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n{time.strftime('%Y年%m月%d日')}")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        
        # 添加分页符开始正文
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        # 添加目录（仅适用于全文）
        if content_type != "outline":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("目 录")
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run.font.bold = True
            
            doc.add_paragraph("\n")
            
            # 使用字段代码添加目录（这是一个占位符，Word会自动填充）
            doc.add_paragraph().add_run("TOC \o \"1-3\" \h \z \\u").font.color.rgb = RGBColor(0, 0, 0)
            
            # 添加分页符开始正文
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        # 定义样式
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 为不同级别标题设置样式
        for i in range(1, 5):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = '黑体'
            heading_style.font.bold = True
            heading_style.font.size = Pt(18 - (i*2))  # 标题1-4级字体大小递减
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(8)
            # 确保标题是黑色的
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            # 设置标题样式和段落格式
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_style.paragraph_format.keep_with_next = True  # 确保标题不会和内容分开
        
        # 解析内容并添加到文档
        if content_type == "outline":
            # 处理大纲内容
            lines = content.split("\n")
            current_level = 0
            is_first_heading = True  # 标记第一个标题
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 清理文本中的*符号
                cleaned_line = clean_markdown_text(line)
                
                # 根据行的格式确定标题级别，并确保移除#符号
                if cleaned_line.startswith("# "):
                    if is_first_heading:
                        # 跳过第一个标题（已经在封面添加过）
                        is_first_heading = False
                        continue
                    title_text = cleaned_line[2:].strip()  # 确保清理干净
                    heading = doc.add_heading(title_text, level=1)
                    # 确保标题文本是黑色的
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif cleaned_line.startswith("## "):
                    title_text = cleaned_line[3:].strip()
                    heading = doc.add_heading(title_text, level=2)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif cleaned_line.startswith("### "):
                    title_text = cleaned_line[4:].strip()
                    heading = doc.add_heading(title_text, level=3)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif cleaned_line.startswith("#### "):
                    title_text = cleaned_line[5:].strip()
                    heading = doc.add_heading(title_text, level=4)
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif cleaned_line.startswith("- "):
                    p = doc.add_paragraph(cleaned_line[2:], style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)
                else:
                    p = doc.add_paragraph(cleaned_line)
                    p.paragraph_format.first_line_indent = Pt(24)
        else:
            # 处理全文内容
            in_paragraph_block = False
            paragraph_text = ""
            
            lines = content.split("\n")
            is_first_heading = True  # 标记第一个标题
            
            for line in lines:
                # 清理文本中的*符号
                cleaned_line = clean_markdown_text(line.strip())
                
                # 处理标题和缩进
                if not cleaned_line:
                    if paragraph_text:
                        # 结束上一个段落块并添加
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Pt(24)
                        p.paragraph_format.line_spacing = 1.5
                        run = p.add_run(paragraph_text)
                        paragraph_text = ""
                    in_paragraph_block = False
                    continue
                
                # 处理标题，确保移除#符号并设置为黑色
                if cleaned_line.startswith("#"):
                    # 如果有缓存的段落文本，先添加
                    if paragraph_text:
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Pt(24)
                        p.paragraph_format.line_spacing = 1.5
                        run = p.add_run(paragraph_text)
                        paragraph_text = ""
                    
                    # 根据行的格式确定标题级别并添加
                    if cleaned_line.startswith("# "):
                        if is_first_heading:
                            # 跳过第一个标题（已经在封面添加过）
                            is_first_heading = False
                            continue
                        title_text = cleaned_line[2:].strip()
                        heading = doc.add_heading(title_text, level=1)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif cleaned_line.startswith("## "):
                        title_text = cleaned_line[3:].strip()
                        heading = doc.add_heading(title_text, level=2)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif cleaned_line.startswith("### "):
                        title_text = cleaned_line[4:].strip()
                        heading = doc.add_heading(title_text, level=3)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif cleaned_line.startswith("#### "):
                        title_text = cleaned_line[5:].strip()
                        heading = doc.add_heading(title_text, level=4)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    in_paragraph_block = False
                
                # 处理列表
                elif cleaned_line.startswith("- "):
                    # 如果有缓存的段落文本，先添加
                    if paragraph_text:
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Pt(24)
                        p.paragraph_format.line_spacing = 1.5
                        run = p.add_run(paragraph_text)
                        paragraph_text = ""
                    
                    p = doc.add_paragraph(cleaned_line[2:], style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)
                    in_paragraph_block = False
                
                # 处理普通段落
                else:
                    # 累积段落文本
                    if not in_paragraph_block:
                        in_paragraph_block = True
                        paragraph_text = cleaned_line
                    else:
                        paragraph_text += " " + cleaned_line
            
            # 添加最后一个可能未处理的段落
            if paragraph_text:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(paragraph_text)
        
        # 添加页码
        section = doc.sections[-1]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 使用正确的页码添加方法
        p.text = "第 "
        run = p.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar)
        run._element.append(instrText)
        run._element.append(fldChar2)
        p.add_run(" 页")
        
        # 生成随机文件名，避免冲突
        filename = f"{uuid.uuid4().hex}.docx"
        file_path = os.path.join("downloads", filename)
        doc.save(file_path)
        
        # 返回文件，使用提取的标题作为文件名
        return FileResponse(
            path=file_path,
            filename=f"{main_title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建Word文档时出错: {str(e)}")


# 定期清理长时间未使用的生成任务
@app.on_event("startup")
async def setup_periodic_cleanup():
    # 这里可以添加定时任务，清理长时间未活动的生成任务
    pass


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True) 